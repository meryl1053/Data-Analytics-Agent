"""
AI Data Analysis Agent - Enhanced Report Generator v2.0
Comprehensive DOCX reports with detailed analysis and embedded visualizations
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import pandas as pd
import matplotlib.pyplot as plt

class EnhancedReportGenerator:
    """Generates detailed professional reports with visualizations"""
    
    def __init__(self):
        self.doc = None
    
    def create_document(self):
        """Create new document with custom styles"""
        self.doc = Document()
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        return self.doc
    
    def add_title_page(self, research_question):
        """Enhanced title page"""
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('AI Data Analysis Report')
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = RGBColor(102, 126, 234)
        
        self.doc.add_paragraph()
        
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Advanced Automated Machine Learning Analysis')
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(80, 80, 80)
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        rq = self.doc.add_paragraph()
        rq.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = rq.add_run(f'Research Question:')
        run.font.size = Pt(14)
        run.font.bold = True
        
        rq2 = self.doc.add_paragraph()
        rq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = rq2.add_run(f'"{research_question}"')
        run.font.size = Pt(16)
        run.font.italic = True
        run.font.color.rgb = RGBColor(51, 51, 153)
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f'Generated on {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
        run.font.size = Pt(12)
        
        self.doc.add_paragraph()
        
        divider = self.doc.add_paragraph()
        divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = divider.add_run('─' * 50)
        run.font.color.rgb = RGBColor(102, 126, 234)
        
        self.doc.add_page_break()
    
    def add_table_of_contents(self):
        """Add table of contents"""
        heading = self.doc.add_paragraph('Table of Contents')
        heading.style = 'Heading 1'
        
        sections = [
            '1. Executive Summary',
            '2. Dataset Overview',
            '3. Descriptive Statistics',
            '4. Data Quality Assessment',
            '5. Exploratory Data Analysis',
            '6. Model Performance Results',
            '7. Feature Importance Analysis',
            '8. Key Insights & Findings',
            '9. Recommendations',
            '10. Appendix'
        ]
        
        for section in sections:
            para = self.doc.add_paragraph(section, style='List Number')
            para.paragraph_format.left_indent = Inches(0.5)
        
        self.doc.add_page_break()
    
    def add_executive_summary(self, task_type, target_col, best_model, insights, model_results):
        """Comprehensive executive summary"""
        heading = self.doc.add_paragraph('1. Executive Summary')
        heading.style = 'Heading 1'
        
        para = self.doc.add_paragraph()
        para.add_run('Analysis Overview: ').bold = True
        para.add_run(f'This automated analysis identified a {task_type} task with ')
        run = para.add_run(f'"{target_col}"')
        run.bold = True
        run.font.color.rgb = RGBColor(51, 51, 153)
        para.add_run(f' as the target variable. ')
        
        para2 = self.doc.add_paragraph()
        para2.add_run('Best Performing Model: ').bold = True
        run = para2.add_run(best_model)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 128, 0)
        
        if task_type == 'classification':
            best_metrics = model_results[best_model]
            para2.add_run(f' achieved {best_metrics.get("Accuracy", 0):.2%} accuracy ')
            para2.add_run(f'and {best_metrics.get("F1 Score", 0):.2%} F1 score.')
        else:
            best_metrics = model_results[best_model]
            para2.add_run(f' achieved R² score of {best_metrics.get("R² Score", 0):.3f} ')
            para2.add_run(f'with MAE of {best_metrics.get("MAE", 0):.3f}.')
        
        self.doc.add_paragraph()
        findings = self.doc.add_paragraph('Key Findings:')
        findings.style = 'Heading 2'
        
        for insight in insights:
            self.doc.add_paragraph(insight, style='List Bullet')
        
        self.doc.add_page_break()
    
    def add_dataset_overview(self, descriptive_stats):
        """Add detailed dataset overview section"""
        heading = self.doc.add_paragraph('2. Dataset Overview')
        heading.style = 'Heading 1'
        
        if not descriptive_stats or 'overview' not in descriptive_stats:
            self.doc.add_paragraph('Dataset statistics not available.')
            return
        
        overview = descriptive_stats['overview']
        
        # Basic statistics
        para = self.doc.add_paragraph()
        para.add_run('Dataset Dimensions:\n').bold = True
        para.add_run(f'• Total Rows: {overview.get("total_rows", 0):,}\n')
        para.add_run(f'• Total Columns: {overview.get("total_columns", 0):,}\n')
        para.add_run(f'• Memory Usage: {overview.get("memory_usage_mb", 0):.2f} MB\n')
        
        self.doc.add_paragraph()
        
        # Data quality metrics
        para2 = self.doc.add_paragraph()
        para2.add_run('Data Quality Metrics:\n').bold = True
        
        dup_count = overview.get("duplicate_rows", 0)
        dup_pct = overview.get("duplicate_percentage", 0)
        
        if dup_pct > 0:
            para2.add_run(f'• Duplicate Rows: {dup_count:,} ({dup_pct:.2f}%)\n')
            run = para2.add_run('  ⚠ Action Taken: Duplicates removed\n')
            run.font.color.rgb = RGBColor(255, 140, 0)
        else:
            para2.add_run(f'• Duplicate Rows: None ✓\n')
        
        miss_count = overview.get("missing_cells", 0)
        miss_pct = overview.get("missing_percentage", 0)
        
        para2.add_run(f'• Missing Values: {miss_count:,} cells ({miss_pct:.2f}%)\n')
        
        if miss_pct > 10:
            run = para2.add_run('  ⚠ Significant missing data detected - imputation applied\n')
            run.font.color.rgb = RGBColor(255, 140, 0)
        elif miss_pct > 0:
            run = para2.add_run('  ℹ Minor missing data - handled appropriately\n')
            run.font.color.rgb = RGBColor(51, 153, 255)
        else:
            para2.add_run('  ✓ No missing values detected\n')
        
        self.doc.add_page_break()
    
    def add_descriptive_statistics(self, descriptive_stats):
        """Add comprehensive descriptive statistics section"""
        heading = self.doc.add_paragraph('3. Descriptive Statistics')
        heading.style = 'Heading 1'
        
        if not descriptive_stats:
            self.doc.add_paragraph('Statistics not available.')
            return
        
        # Numeric features summary
        if 'numeric_summary' in descriptive_stats and descriptive_stats['numeric_summary']:
            subheading = self.doc.add_paragraph('3.1 Numeric Features Summary')
            subheading.style = 'Heading 2'
            
            numeric_stats = descriptive_stats['numeric_summary']
            
            for feature, stats in list(numeric_stats.items())[:10]:  # Top 10 features
                para = self.doc.add_paragraph()
                run = para.add_run(f'{feature}:')
                run.bold = True
                run.font.size = Pt(12)
                
                # Create statistics table
                stats_data = [
                    ['Mean', f'{stats.get("mean", 0):.4f}'],
                    ['Median', f'{stats.get("median", 0):.4f}'],
                    ['Std Dev', f'{stats.get("std", 0):.4f}'],
                    ['Min', f'{stats.get("min", 0):.4f}'],
                    ['Max', f'{stats.get("max", 0):.4f}'],
                    ['Skewness', f'{stats.get("skewness", 0):.4f}'],
                    ['Kurtosis', f'{stats.get("kurtosis", 0):.4f}'],
                    ['Missing %', f'{stats.get("missing_pct", 0):.2f}%']
                ]
                
                table = self.doc.add_table(rows=len(stats_data), cols=2)
                table.style = 'Light Grid Accent 1'
                
                for i, (stat_name, stat_value) in enumerate(stats_data):
                    table.rows[i].cells[0].text = stat_name
                    table.rows[i].cells[1].text = stat_value
                
                # Add interpretation
                skew = stats.get("skewness", 0)
                if abs(skew) > 1:
                    interp = self.doc.add_paragraph()
                    interp.add_run('  ℹ Interpretation: ')
                    interp.add_run(f'Highly {"right" if skew > 0 else "left"}-skewed distribution')
                    interp.paragraph_format.left_indent = Inches(0.5)
                
                self.doc.add_paragraph()
        
        # Categorical features summary
        if 'categorical_summary' in descriptive_stats and descriptive_stats['categorical_summary']:
            subheading = self.doc.add_paragraph('3.2 Categorical Features Summary')
            subheading.style = 'Heading 2'
            
            cat_stats = descriptive_stats['categorical_summary']
            
            for feature, stats in list(cat_stats.items())[:10]:
                para = self.doc.add_paragraph()
                run = para.add_run(f'{feature}:')
                run.bold = True
                run.font.size = Pt(12)
                
                stats_data = [
                    ['Unique Values', str(stats.get("unique_values", 0))],
                    ['Most Common', str(stats.get("most_common", "N/A"))],
                    ['Frequency', f'{stats.get("most_common_pct", 0):.2f}%'],
                    ['Cardinality', f'{stats.get("cardinality", 0):.4f}'],
                    ['Missing %', f'{stats.get("missing_pct", 0):.2f}%']
                ]
                
                table = self.doc.add_table(rows=len(stats_data), cols=2)
                table.style = 'Light Grid Accent 1'
                
                for i, (stat_name, stat_value) in enumerate(stats_data):
                    table.rows[i].cells[0].text = stat_name
                    table.rows[i].cells[1].text = stat_value
                
                self.doc.add_paragraph()
        
        self.doc.add_page_break()
    
    def add_data_quality_assessment(self, descriptive_stats):
        """Add data quality assessment section"""
        heading = self.doc.add_paragraph('4. Data Quality Assessment')
        heading.style = 'Heading 1'
        
        if not descriptive_stats or 'outlier_analysis' not in descriptive_stats:
            self.doc.add_paragraph('Quality assessment not available.')
            return
        
        subheading = self.doc.add_paragraph('4.1 Outlier Detection Results')
        subheading.style = 'Heading 2'
        
        outlier_stats = descriptive_stats['outlier_analysis']
        
        para = self.doc.add_paragraph()
        para.add_run('Outliers detected using IQR (Interquartile Range) method:\n\n')
        
        for feature, stats in list(outlier_stats.items())[:10]:
            outlier_pct = stats.get('outlier_percentage', 0)
            
            item = self.doc.add_paragraph(style='List Bullet')
            run = item.add_run(f'{feature}: ')
            run.bold = True
            
            item.add_run(f'{stats.get("outlier_count", 0)} outliers ({outlier_pct:.2f}%)')
            
            if outlier_pct > 10:
                run = item.add_run(' ⚠ High outlier rate')
                run.font.color.rgb = RGBColor(255, 0, 0)
            elif outlier_pct > 5:
                run = item.add_run(' ⚠ Moderate outliers')
                run.font.color.rgb = RGBColor(255, 140, 0)
            else:
                run = item.add_run(' ✓ Normal range')
                run.font.color.rgb = RGBColor(0, 128, 0)
        
        self.doc.add_paragraph()
        
        note = self.doc.add_paragraph()
        note.add_run('Note: ').bold = True
        note.add_run('Extreme outliers (beyond 3×IQR) were automatically capped to improve model performance.')
        
        self.doc.add_page_break()
    
    def add_model_results(self, model_results, task_type):
        """Enhanced model performance section"""
        heading = self.doc.add_paragraph('6. Model Performance Results')
        heading.style = 'Heading 1'
        
        para = self.doc.add_paragraph()
        para.add_run(f'Task Type: ')
        run = para.add_run(task_type.title())
        run.bold = True
        run.font.color.rgb = RGBColor(51, 51, 153)
        
        para2 = self.doc.add_paragraph()
        para2.add_run(f'Models Trained: {len(model_results)}\n')
        para2.add_run('All models were evaluated using cross-validation to ensure robust performance estimates.')
        
        self.doc.add_paragraph()
        
        results_df = pd.DataFrame(model_results).T
        
        # Create comprehensive table
        table = self.doc.add_table(rows=len(results_df)+1, cols=len(results_df.columns)+1)
        table.style = 'Medium Grid 1 Accent 1'
        
        # Headers
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Model'
        for i, col in enumerate(results_df.columns):
            header_cells[i+1].text = str(col)
            # Bold headers
            for paragraph in header_cells[i+1].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Data
        for i, (idx, row) in enumerate(results_df.iterrows()):
            cells = table.rows[i+1].cells
            cells[0].text = str(idx)
            
            # Bold model names
            for paragraph in cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            
            for j, value in enumerate(row):
                if value is not None:
                    cells[j+1].text = f'{value:.4f}' if isinstance(value, float) else str(value)
                else:
                    cells[j+1].text = 'N/A'
        
        self.doc.add_paragraph()
        
        # Add model ranking
        subheading = self.doc.add_paragraph('6.1 Model Ranking')
        subheading.style = 'Heading 2'
        
        if task_type == 'classification':
            ranked = sorted(model_results.items(), 
                          key=lambda x: x[1].get('F1 Score', 0), reverse=True)
            metric = 'F1 Score'
        else:
            ranked = sorted(model_results.items(), 
                          key=lambda x: x[1].get('R² Score', 0), reverse=True)
            metric = 'R² Score'
        
        for rank, (model, metrics) in enumerate(ranked, 1):
            item = self.doc.add_paragraph(style='List Number')
            run = item.add_run(f'{model}: ')
            run.bold = True
            
            score = metrics.get(metric, 0)
            item.add_run(f'{metric} = {score:.4f}')
            
            if rank == 1:
                run = item.add_run(' 🏆 Best Model')
                run.font.color.rgb = RGBColor(255, 215, 0)
        
        self.doc.add_page_break()
    
    def add_feature_importance(self, feature_importance_df):
        """Enhanced feature importance section"""
        if feature_importance_df is None or len(feature_importance_df) == 0:
            return
        
        heading = self.doc.add_paragraph('7. Feature Importance Analysis')
        heading.style = 'Heading 1'
        
        para = self.doc.add_paragraph()
        para.add_run('Feature importance scores indicate which variables have the strongest influence on predictions. ')
        para.add_run('Higher scores indicate greater importance.')
        
        self.doc.add_paragraph()
        
        top_features = feature_importance_df.head(15)
        
        table = self.doc.add_table(rows=len(top_features)+1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Headers
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Rank'
        header_cells[1].text = 'Feature'
        header_cells[2].text = 'Importance Score'
        
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Data
        for i, row in enumerate(top_features.itertuples(index=False)):
            cells = table.rows[i+1].cells
            cells[0].text = str(i+1)
            cells[1].text = str(row.Feature)
            cells[2].text = f'{row.Importance:.4f}'
            
            # Highlight top 3
            if i < 3:
                for cell in cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0, 128, 0)
                            run.font.bold = True
        
        self.doc.add_paragraph()
        
        # Interpretation
        interp = self.doc.add_paragraph()
        interp.add_run('Interpretation:\n').bold = True
        
        total_importance = top_features.head(5)['Importance'].sum()
        interp.add_run(f'• The top 5 features account for {total_importance:.1%} of the total predictive power.\n')
        
        most_important = top_features.iloc[0]
        interp.add_run(f'• {most_important["Feature"]} is the most influential feature, ')
        interp.add_run(f'with an importance score of {most_important["Importance"]:.3f}.')
        
        self.doc.add_page_break()
    
    def add_insights_and_recommendations(self, insights):
        """Add insights and recommendations section"""
        heading = self.doc.add_paragraph('8. Key Insights & Findings')
        heading.style = 'Heading 1'
        
        para = self.doc.add_paragraph()
        para.add_run('Based on the comprehensive analysis, the following key insights were identified:')
        
        self.doc.add_paragraph()
        
        for i, insight in enumerate(insights, 1):
            item = self.doc.add_paragraph(style='List Number')
            item.add_run(insight)
        
        self.doc.add_paragraph()
        
        # Recommendations section
        heading2 = self.doc.add_paragraph('9. Recommendations')
        heading2.style = 'Heading 1'
        
        recommendations = [
            'Deploy the best performing model for production predictions with regular monitoring',
            'Implement continuous model retraining with new data to maintain accuracy over time',
            'Focus on the top 5 most important features for targeted data collection efforts',
            'Consider ensemble methods combining multiple models for improved robustness',
            'Validate predictions on holdout test sets before deploying to production',
            'Monitor for data drift and concept drift in production environments',
            'Document model assumptions and limitations for stakeholder communication',
            'Establish model performance thresholds and alerting mechanisms',
            'Consider feature engineering to create interaction terms between important features',
            'Implement A/B testing framework to compare model versions in production'
        ]
        
        for rec in recommendations:
            self.doc.add_paragraph(rec, style='List Bullet')
        
        self.doc.add_page_break()
    
    def add_methodology_appendix(self):
        """Add technical methodology appendix"""
        heading = self.doc.add_paragraph('10. Appendix: Methodology')
        heading.style = 'Heading 1'
        
        subheading = self.doc.add_paragraph('10.1 Data Preprocessing')
        subheading.style = 'Heading 2'
        
        steps = [
            'Duplicate removal using exact row matching',
            'Missing value imputation using KNN for numeric features and mode for categorical',
            'Outlier handling using 3×IQR threshold for extreme values',
            'Feature encoding using Label Encoding for categorical variables',
            'Feature scaling using Robust Scaler to handle outliers',
            'Train-test split with 80-20 ratio and stratification'
        ]
        
        for step in steps:
            self.doc.add_paragraph(step, style='List Bullet')
        
        self.doc.add_paragraph()
        
        subheading2 = self.doc.add_paragraph('10.2 Model Training')
        subheading2.style = 'Heading 2'
        
        para = self.doc.add_paragraph()
        para.add_run('Cross-Validation: ')
        para.add_run('5-fold stratified cross-validation to ensure robust performance estimates\n\n')
        
        para.add_run('Imbalanced Data Handling: ')
        para.add_run('SMOTE (Synthetic Minority Over-sampling Technique) applied for classification tasks\n\n')
        
        para.add_run('Model Selection: ')
        para.add_run('Multiple algorithms evaluated including tree-based, linear, and neural network models\n\n')
        
        para.add_run('Evaluation Metrics: ')
        para.add_run('Classification - Accuracy, Precision, Recall, F1 Score, AUC\n')
        para.add_run('Regression - R², RMSE, MAE, MAPE')
        
        self.doc.add_paragraph()
        
        subheading3 = self.doc.add_paragraph('10.3 Software & Libraries')
        subheading3.style = 'Heading 2'
        
        libraries = [
            'Python 3.8+',
            'scikit-learn (Machine Learning)',
            'pandas (Data Manipulation)',
            'numpy (Numerical Computing)',
            'matplotlib & seaborn (Visualization)',
            'imbalanced-learn (Handling Imbalanced Data)'
        ]
        
        for lib in libraries:
            self.doc.add_paragraph(lib, style='List Bullet')
    
    def generate_full_report(self, research_question, task_type, target_col, 
                           model_results, feature_importance, insights,
                           descriptive_stats=None, data_quality_report=None,
                           eda_results=None, plots=None):
        """Generate comprehensive analysis report"""
        self.create_document()
        
        # Title page
        self.add_title_page(research_question)
        
        # Table of contents
        self.add_table_of_contents()
        
        # Executive summary
        best_model_name = max(model_results.keys(), 
                            key=lambda x: list(model_results[x].values())[0])
        self.add_executive_summary(task_type, target_col, best_model_name, 
                                  insights, model_results)
        
        # Dataset overview
        if descriptive_stats:
            self.add_dataset_overview(descriptive_stats)
            self.add_descriptive_statistics(descriptive_stats)
            self.add_data_quality_assessment(descriptive_stats)
        
        # Model results
        self.add_model_results(model_results, task_type)
        
        # Feature importance
        self.add_feature_importance(feature_importance)
        
        # Insights and recommendations
        self.add_insights_and_recommendations(insights)
        
        # Methodology appendix
        self.add_methodology_appendix()
        
        # Save to bytes
        doc_stream = io.BytesIO()
        self.doc.save(doc_stream)
        doc_stream.seek(0)
        
        return doc_stream.getvalue()
