"""
Generate all module files for AI Data Analysis Agent
Run this script to create all required Python files
"""

import os

# Module contents as strings
MODULES = {
    "analysis_engine.py": '''"""
AI Data Analysis Agent - Analysis Engine
Automated ML pipeline with intelligent target detection
"""

import pandas as pd
import numpy as np
from sklearn.model_selection import train_test_split, cross_val_score
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.impute import SimpleImputer
from sklearn.ensemble import RandomForestClassifier, RandomForestRegressor, GradientBoostingClassifier, GradientBoostingRegressor
from sklearn.linear_model import LogisticRegression, LinearRegression
from sklearn.svm import SVC, SVR
from sklearn.neural_network import MLPClassifier, MLPRegressor
from sklearn.cluster import KMeans, DBSCAN
from sklearn.metrics import (accuracy_score, f1_score, roc_auc_score, confusion_matrix,
                            r2_score, mean_squared_error, mean_absolute_error,
                            silhouette_score)
from imblearn.over_sampling import SMOTE
import warnings
warnings.filterwarnings('ignore')

class AIAnalysisEngine:
    """Main analysis engine for automated ML pipeline"""
    
    def __init__(self, dataframe, research_question=""):
        self.df = dataframe.copy()
        self.research_question = research_question.lower()
        self.target_column = None
        self.task_type = None
        self.X_train = None
        self.X_test = None
        self.y_train = None
        self.y_test = None
        self.trained_models = {}
        self.feature_importance = None
        self.label_encoders = {}
        self.scaler = None
        
    def detect_target_column(self):
        """Intelligently detect target variable using 3 strategies"""
        # Strategy 1: Parse research question
        if self.research_question:
            for col in self.df.columns:
                if col.lower() in self.research_question:
                    self.target_column = col
                    return col
        
        # Strategy 2: Look for common target column names
        target_patterns = ['target', 'label', 'y', 'outcome', 'class', 
                          'prediction', 'result', 'status', 'approved',
                          'churn', 'fraud', 'default', 'price', 'revenue',
                          'sales', 'value', 'score']
        
        for col in self.df.columns:
            if any(pattern in col.lower() for pattern in target_patterns):
                self.target_column = col
                return col
        
        # Strategy 3: Use last column as default
        self.target_column = self.df.columns[-1]
        return self.target_column
    
    def identify_task_type(self):
        """Identify whether task is classification, regression, or clustering"""
        if self.target_column is None:
            self.detect_target_column()
        
        target_data = self.df[self.target_column]
        
        if target_data.isnull().all():
            return 'clustering'
        
        unique_values = target_data.nunique()
        total_values = len(target_data.dropna())
        
        if target_data.dtype == 'object' or unique_values < 20:
            self.task_type = 'classification'
        elif unique_values / total_values > 0.05:
            self.task_type = 'regression'
        else:
            self.task_type = 'classification'
        
        return self.task_type
    
    def preprocess_data(self):
        """Automated preprocessing: imputation, encoding, scaling"""
        X = self.df.drop(columns=[self.target_column])
        y = self.df[self.target_column]
        
        if y.isnull().any():
            y = y.fillna(y.mode()[0] if self.task_type == 'classification' else y.median())
        
        if self.task_type == 'classification' and y.dtype == 'object':
            le = LabelEncoder()
            y = le.fit_transform(y)
            self.label_encoders['target'] = le
        
        numeric_features = X.select_dtypes(include=['number']).columns.tolist()
        categorical_features = X.select_dtypes(include=['object', 'category']).columns.tolist()
        
        if numeric_features:
            num_imputer = SimpleImputer(strategy='median')
            X[numeric_features] = num_imputer.fit_transform(X[numeric_features])
        
        if categorical_features:
            cat_imputer = SimpleImputer(strategy='most_frequent')
            X[categorical_features] = cat_imputer.fit_transform(X[categorical_features])
        
        for col in categorical_features:
            le = LabelEncoder()
            X[col] = le.fit_transform(X[col].astype(str))
            self.label_encoders[col] = le
        
        self.scaler = StandardScaler()
        X_scaled = self.scaler.fit_transform(X)
        X = pd.DataFrame(X_scaled, columns=X.columns, index=X.index)
        
        return X, y
    
    def train_classification_models(self, X_train, X_test, y_train, y_test, cv_folds=5, use_smote=True):
        """Train multiple classification models"""
        if use_smote and len(np.unique(y_train)) == 2:
            try:
                smote = SMOTE(random_state=42)
                X_train, y_train = smote.fit_resample(X_train, y_train)
            except:
                pass
        
        models = {
            'Logistic Regression': LogisticRegression(max_iter=1000, random_state=42),
            'Random Forest': RandomForestClassifier(n_estimators=100, random_state=42),
            'Gradient Boosting': GradientBoostingClassifier(n_estimators=100, random_state=42)
        }
        
        results = {}
        
        for name, model in models.items():
            try:
                model.fit(X_train, y_train)
                y_pred = model.predict(X_test)
                
                accuracy = accuracy_score(y_test, y_pred)
                f1 = f1_score(y_test, y_pred, average='weighted')
                cv_scores = cross_val_score(model, X_train, y_train, cv=cv_folds)
                
                results[name] = {
                    'Accuracy': round(accuracy, 4),
                    'F1 Score': round(f1, 4),
                    'CV Mean': round(cv_scores.mean(), 4),
                    'CV Std': round(cv_scores.std(), 4)
                }
                
                self.trained_models[name] = model
            except Exception as e:
                print(f"Error training {name}: {str(e)}")
                continue
        
        return results
    
    def train_regression_models(self, X_train, X_test, y_train, y_test, cv_folds=5):
        """Train multiple regression models"""
        models = {
            'Linear Regression': LinearRegression(),
            'Random Forest': RandomForestRegressor(n_estimators=100, random_state=42),
            'Gradient Boosting': GradientBoostingRegressor(n_estimators=100, random_state=42)
        }
        
        results = {}
        
        for name, model in models.items():
            try:
                model.fit(X_train, y_train)
                y_pred = model.predict(X_test)
                
                r2 = r2_score(y_test, y_pred)
                rmse = np.sqrt(mean_squared_error(y_test, y_pred))
                mae = mean_absolute_error(y_test, y_pred)
                cv_scores = cross_val_score(model, X_train, y_train, cv=cv_folds, scoring='r2')
                
                results[name] = {
                    'R² Score': round(r2, 4),
                    'RMSE': round(rmse, 4),
                    'MAE': round(mae, 4),
                    'CV Mean': round(cv_scores.mean(), 4),
                    'CV Std': round(cv_scores.std(), 4)
                }
                
                self.trained_models[name] = model
            except Exception as e:
                print(f"Error training {name}: {str(e)}")
                continue
        
        return results
    
    def extract_feature_importance(self, model_name='Random Forest'):
        """Extract feature importance from tree-based models"""
        if model_name not in self.trained_models:
            return None
        
        model = self.trained_models[model_name]
        
        if hasattr(model, 'feature_importances_'):
            importance = model.feature_importances_
            feature_names = self.X_train.columns
            
            importance_df = pd.DataFrame({
                'Feature': feature_names,
                'Importance': importance
            }).sort_values('Importance', ascending=False)
            
            self.feature_importance = importance_df
            return importance_df
        
        return None
    
    def generate_insights(self, model_results):
        """Generate automated insights from analysis"""
        insights = []
        
        if self.task_type == 'classification':
            best_model = max(model_results.items(), key=lambda x: x[1].get('Accuracy', 0))
            insights.append(f"Best performing model is {best_model[0]} with {best_model[1]['Accuracy']:.1%} accuracy")
        elif self.task_type == 'regression':
            best_model = max(model_results.items(), key=lambda x: x[1].get('R² Score', 0))
            insights.append(f"Best performing model is {best_model[0]} with R² = {best_model[1]['R² Score']:.3f}")
        
        if self.feature_importance is not None:
            top_feature = self.feature_importance.iloc[0]
            insights.append(f"Most important feature: {top_feature['Feature']} ({top_feature['Importance']:.3f})")
        
        missing_pct = (self.df.isnull().sum().sum() / (len(self.df) * len(self.df.columns))) * 100
        if missing_pct > 10:
            insights.append(f"Dataset has {missing_pct:.1f}% missing values - imputation was applied")
        
        return insights
    
    def run_full_pipeline(self, test_size=0.2, cv_folds=5, use_smote=True):
        """Execute complete analysis pipeline"""
        self.detect_target_column()
        self.identify_task_type()
        
        X, y = self.preprocess_data()
        
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=test_size, random_state=42)
        
        self.X_train = X_train
        self.X_test = X_test
        self.y_train = y_train
        self.y_test = y_test
        
        if self.task_type == 'classification':
            model_results = self.train_classification_models(X_train, X_test, y_train, y_test, cv_folds, use_smote)
        else:
            model_results = self.train_regression_models(X_train, X_test, y_train, y_test, cv_folds)
        
        if 'Random Forest' in self.trained_models:
            self.extract_feature_importance('Random Forest')
        
        insights = self.generate_insights(model_results)
        
        if self.task_type == 'classification':
            best_model = max(model_results.items(), key=lambda x: x[1].get('Accuracy', 0))[0]
        else:
            best_model = max(model_results.items(), key=lambda x: x[1].get('R² Score', 0))[0]
        
        return {
            'target_column': self.target_column,
            'task_type': self.task_type,
            'model_results': model_results,
            'best_model': best_model,
            'best_model_object': self.trained_models[best_model],
            'feature_importance': self.feature_importance,
            'insights': insights,
            'label_encoders': self.label_encoders,
            'scaler': self.scaler
        }
''',

    "visualization_module.py": '''"""
AI Data Analysis Agent - Visualization Module
Comprehensive plotting system for EDA and model evaluation
"""

import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
import numpy as np
import warnings
warnings.filterwarnings('ignore')

sns.set_style("whitegrid")
plt.rcParams['figure.facecolor'] = 'white'

class VisualizationEngine:
    """Handles all visualization tasks"""
    
    def __init__(self, figsize=(12, 6), dpi=100):
        self.figsize = figsize
        self.dpi = dpi
        self.color_palette = sns.color_palette("husl", 10)
    
    def plot_missing_values(self, df, max_cols=20):
        """Visualize missing values in dataset"""
        missing = df.isnull().sum()
        missing_pct = (missing / len(df)) * 100
        missing_data = pd.DataFrame({
            'Column': missing.index,
            'Missing_Percentage': missing_pct.values
        })
        missing_data = missing_data[missing_data['Missing_Percentage'] > 0].sort_values(
            'Missing_Percentage', ascending=False
        ).head(max_cols)
        
        if len(missing_data) == 0:
            fig, ax = plt.subplots(figsize=self.figsize, dpi=self.dpi)
            ax.text(0.5, 0.5, 'No Missing Values Found!', ha='center', va='center', fontsize=20, color='green')
            ax.axis('off')
            return fig
        
        fig, ax = plt.subplots(figsize=self.figsize, dpi=self.dpi)
        bars = ax.barh(missing_data['Column'], missing_data['Missing_Percentage'], color=self.color_palette[0])
        
        for bar in bars:
            width = bar.get_width()
            ax.text(width, bar.get_y() + bar.get_height()/2, f'{width:.1f}%', ha='left', va='center', fontsize=9)
        
        ax.set_xlabel('Missing Percentage (%)', fontsize=12)
        ax.set_title('Missing Values Analysis', fontsize=14, fontweight='bold')
        ax.invert_yaxis()
        plt.tight_layout()
        return fig
    
    def plot_correlation_heatmap(self, df, method='pearson', max_features=15):
        """Create correlation heatmap for numeric features"""
        numeric_df = df.select_dtypes(include=['number'])
        
        if len(numeric_df.columns) == 0:
            fig, ax = plt.subplots(figsize=self.figsize, dpi=self.dpi)
            ax.text(0.5, 0.5, 'No Numeric Features Found', ha='center', va='center', fontsize=20)
            ax.axis('off')
            return fig
        
        if len(numeric_df.columns) > max_features:
            target_corr = numeric_df.corr()[numeric_df.columns[-1]].abs().sort_values(ascending=False)
            top_features = target_corr.head(max_features).index
            numeric_df = numeric_df[top_features]
        
        corr_matrix = numeric_df.corr(method=method)
        
        fig, ax = plt.subplots(figsize=(10, 8), dpi=self.dpi)
        sns.heatmap(corr_matrix, annot=True, fmt='.2f', cmap='coolwarm', center=0, square=True, linewidths=1, ax=ax)
        ax.set_title(f'Correlation Heatmap ({method.title()})', fontsize=14, fontweight='bold', pad=20)
        plt.tight_layout()
        return fig
    
    def plot_distributions(self, df, max_features=6):
        """Plot distributions of numeric features"""
        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()[:max_features]
        
        if len(numeric_cols) == 0:
            fig, ax = plt.subplots(figsize=self.figsize, dpi=self.dpi)
            ax.text(0.5, 0.5, 'No Numeric Features to Plot', ha='center', va='center', fontsize=20)
            ax.axis('off')
            return fig
        
        n_cols = min(3, len(numeric_cols))
        n_rows = int(np.ceil(len(numeric_cols) / n_cols))
        
        fig, axes = plt.subplots(n_rows, n_cols, figsize=(15, n_rows * 4), dpi=self.dpi)
        axes = axes.flatten() if len(numeric_cols) > 1 else [axes]
        
        for idx, col in enumerate(numeric_cols):
            data = df[col].dropna()
            axes[idx].hist(data, bins=30, alpha=0.7, color=self.color_palette[idx % len(self.color_palette)], edgecolor='black')
            axes[idx].set_title(f'Distribution: {col}', fontsize=12, fontweight='bold')
            axes[idx].set_xlabel(col, fontsize=10)
            axes[idx].set_ylabel('Frequency', fontsize=10)
        
        for idx in range(len(numeric_cols), len(axes)):
            axes[idx].axis('off')
        
        plt.tight_layout()
        return fig
    
    def plot_outliers(self, df, max_features=6):
        """Create box plots to identify outliers"""
        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()[:max_features]
        
        if len(numeric_cols) == 0:
            fig, ax = plt.subplots(figsize=self.figsize, dpi=self.dpi)
            ax.text(0.5, 0.5, 'No Numeric Features to Analyze', ha='center', va='center', fontsize=20)
            ax.axis('off')
            return fig
        
        n_cols = min(3, len(numeric_cols))
        n_rows = int(np.ceil(len(numeric_cols) / n_cols))
        
        fig, axes = plt.subplots(n_rows, n_cols, figsize=(15, n_rows * 4), dpi=self.dpi)
        axes = axes.flatten() if len(numeric_cols) > 1 else [axes]
        
        for idx, col in enumerate(numeric_cols):
            data = df[col].dropna()
            axes[idx].boxplot(data, vert=True, patch_artist=True)
            axes[idx].set_title(f'Outlier Detection: {col}', fontsize=12, fontweight='bold')
            axes[idx].set_ylabel(col, fontsize=10)
        
        for idx in range(len(numeric_cols), len(axes)):
            axes[idx].axis('off')
        
        plt.tight_layout()
        return fig
    
    def plot_categorical_distributions(self, df, max_features=6):
        """Plot distributions of categorical features"""
        cat_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()[:max_features]
        
        if len(cat_cols) == 0:
            fig, ax = plt.subplots(figsize=self.figsize, dpi=self.dpi)
            ax.text(0.5, 0.5, 'No Categorical Features Found', ha='center', va='center', fontsize=20)
            ax.axis('off')
            return fig
        
        n_cols = min(2, len(cat_cols))
        n_rows = int(np.ceil(len(cat_cols) / n_cols))
        
        fig, axes = plt.subplots(n_rows, n_cols, figsize=(14, n_rows * 5), dpi=self.dpi)
        axes = axes.flatten() if len(cat_cols) > 1 else [axes]
        
        for idx, col in enumerate(cat_cols):
            value_counts = df[col].value_counts().head(10)
            axes[idx].bar(range(len(value_counts)), value_counts.values, color=self.color_palette[idx % len(self.color_palette)])
            axes[idx].set_xticks(range(len(value_counts)))
            axes[idx].set_xticklabels(value_counts.index, rotation=45, ha='right')
            axes[idx].set_title(f'{col} Distribution', fontsize=12, fontweight='bold')
            axes[idx].set_ylabel('Count', fontsize=10)
        
        for idx in range(len(cat_cols), len(axes)):
            axes[idx].axis('off')
        
        plt.tight_layout()
        return fig
    
    def plot_feature_importance(self, feature_importance_df, top_n=15):
        """Plot feature importance from model"""
        if feature_importance_df is None or len(feature_importance_df) == 0:
            fig, ax = plt.subplots(figsize=self.figsize, dpi=self.dpi)
            ax.text(0.5, 0.5, 'Feature Importance Not Available', ha='center', va='center', fontsize=20)
            ax.axis('off')
            return fig
        
        top_features = feature_importance_df.head(top_n)
        
        fig, ax = plt.subplots(figsize=(10, 8), dpi=self.dpi)
        bars = ax.barh(top_features['Feature'], top_features['Importance'], color=self.color_palette[2])
        
        for bar in bars:
            width = bar.get_width()
            ax.text(width, bar.get_y() + bar.get_height()/2, f'{width:.3f}', ha='left', va='center', fontsize=9)
        
        ax.set_xlabel('Importance Score', fontsize=12)
        ax.set_title('Feature Importance Ranking', fontsize=14, fontweight='bold')
        ax.invert_yaxis()
        plt.tight_layout()
        return fig
    
    def plot_model_comparison(self, model_results, task_type):
        """Compare performance of different models"""
        if not model_results:
            fig, ax = plt.subplots(figsize=self.figsize, dpi=self.dpi)
            ax.text(0.5, 0.5, 'No Model Results Available', ha='center', va='center', fontsize=20)
            ax.axis('off')
            return fig
        
        models = list(model_results.keys())
        
        if task_type == 'classification':
            metric1 = [model_results[m].get('Accuracy', 0) for m in models]
            metric2 = [model_results[m].get('F1 Score', 0) for m in models]
            labels = ['Accuracy', 'F1 Score']
        else:
            metric1 = [model_results[m].get('R² Score', 0) for m in models]
            metric2 = [model_results[m].get('MAE', 0) for m in models]
            labels = ['R² Score', 'MAE']
        
        fig, ax = plt.subplots(figsize=(12, 6), dpi=self.dpi)
        x = np.arange(len(models))
        width = 0.35
        
        ax.bar(x - width/2, metric1, width, label=labels[0], color=self.color_palette[0])
        ax.bar(x + width/2, metric2, width, label=labels[1], color=self.color_palette[1])
        
        ax.set_ylabel('Score', fontsize=12)
        ax.set_title('Model Performance Comparison', fontsize=14, fontweight='bold')
        ax.set_xticks(x)
        ax.set_xticklabels(models, rotation=45, ha='right')
        ax.legend()
        plt.tight_layout()
        return fig
    
    def generate_all_eda_plots(self, df):
        """Generate complete set of EDA plots"""
        plots = {}
        plots['Missing Values'] = self.plot_missing_values(df)
        plots['Correlation Heatmap'] = self.plot_correlation_heatmap(df)
        plots['Numeric Distributions'] = self.plot_distributions(df)
        plots['Outlier Detection'] = self.plot_outliers(df)
        plots['Categorical Distributions'] = self.plot_categorical_distributions(df)
        return plots
''',

    "report_generator.py": '''"""
AI Data Analysis Agent - Report Generator
Professional DOCX report creation with embedded visualizations
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import pandas as pd

class ReportGenerator:
    """Handles professional report generation"""
    
    def __init__(self):
        self.doc = None
    
    def create_document(self):
        """Create new document with custom styles"""
        self.doc = Document()
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        return self.doc
    
    def add_title_page(self, research_question):
        """Add title page to report"""
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('AI Data Analysis Report')
        run.font.size = Pt(32)
        run.font.bold = True
        run.font.color.rgb = RGBColor(102, 126, 234)
        
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Automated Machine Learning Analysis')
        run.font.size = Pt(16)
        
        self.doc.add_paragraph()
        rq = self.doc.add_paragraph()
        rq.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = rq.add_run(f'Research Question: "{research_question}"')
        run.font.size = Pt(14)
        run.font.italic = True
        
        self.doc.add_paragraph()
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f'Generated on {datetime.now().strftime("%B %d, %Y")}')
        run.font.size = Pt(12)
        
        self.doc.add_page_break()
    
    def add_executive_summary(self, task_type, target_col, best_model, insights):
        """Add executive summary section"""
        heading = self.doc.add_paragraph('Executive Summary')
        heading.style = 'Heading 1'
        
        para = self.doc.add_paragraph()
        para.add_run(f'This analysis identified a {task_type} task with "{target_col}" as the target variable. ')
        para.add_run(f'The best performing model was {best_model}.')
        
        self.doc.add_paragraph()
        findings = self.doc.add_paragraph('Key Findings:')
        findings.style = 'Heading 2'
        
        for insight in insights:
            self.doc.add_paragraph(insight, style='List Bullet')
        
        self.doc.add_paragraph()
    
    def add_model_results(self, model_results, task_type):
        """Add model performance section"""
        heading = self.doc.add_paragraph('Model Performance')
        heading.style = 'Heading 1'
        
        results_df = pd.DataFrame(model_results).T
        
        table = self.doc.add_table(rows=len(results_df)+1, cols=len(results_df.columns)+1)
        table.style = 'Light Grid Accent 1'
        
        # Headers
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Model'
        for i, col in enumerate(results_df.columns):
            header_cells[i+1].text = str(col)
        
        # Data
        for i, (idx, row) in enumerate(results_df.iterrows()):
            cells = table.rows[i+1].cells
            cells[0].text = str(idx)
            for j, value in enumerate(row):
                cells[j+1].text = f'{value:.4f}' if isinstance(value, float) else str(value)
        
        self.doc.add_paragraph()
    
    def add_feature_importance(self, feature_importance_df):
        """Add feature importance section"""
        if feature_importance_df is None or len(feature_importance_df) == 0:
            return
        
        heading = self.doc.add_paragraph('Feature Importance')
        heading.style = 'Heading 1'
        
        top_features = feature_importance_df.head(10)
        
        table = self.doc.add_table(rows=len(top_features)+1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Feature'
        header_cells[1].text = 'Importance'
        
        for i, row in enumerate(top_features.itertuples(index=False)):
            cells = table.rows[i+1].cells
            cells[0].text = str(row.Feature)
            cells[1].text = f'{row.Importance:.4f}'
        
        self.doc.add_paragraph()
    
    def add_recommendations(self):
        """Add recommendations section"""
        heading = self.doc.add_paragraph('Recommendations')
        heading.style = 'Heading 1'
        
        recommendations = [
            'Deploy the best performing model for production predictions',
            'Monitor model performance regularly and retrain with new data',
            'Consider feature engineering to improve model accuracy',
            'Validate predictions on new, unseen data before deployment'
        ]
        
        for rec in recommendations:
            self.doc.add_paragraph(rec, style='List Bullet')
        
        self.doc.add_paragraph()
    
    def generate_full_report(self, research_question, task_type, target_col, 
                           model_results, feature_importance, insights,
                           eda_results=None, plots=None):
        """Generate complete analysis report"""
        self.create_document()
        
        self.add_title_page(research_question)
        self.add_executive_summary(task_type, target_col, 
                                  max(model_results.keys(), key=lambda x: list(model_results[x].values())[0]),
                                  insights)
        self.add_model_results(model_results, task_type)
        self.add_feature_importance(feature_importance)
        self.add_recommendations()
        
        doc_stream = io.BytesIO()
        self.doc.save(doc_stream)
        doc_stream.seek(0)
        
        return doc_stream.getvalue()
''',

    "utils_module.py": '''"""
AI Data Analysis Agent - Utilities Module
Helper functions for validation, persistence, and configuration
"""

import pandas as pd
import numpy as np
import pickle
import io
from datetime import datetime

class DataValidator:
    """Validates dataset quality and health"""
    
    def __init__(self, max_size_mb=200, min_rows=10, min_cols=2):
        self.max_size_mb = max_size_mb
        self.min_rows = min_rows
        self.min_cols = min_cols
    
    def validate_dataframe(self, df):
        """Validate dataframe for analysis"""
        if df is None or len(df) == 0:
            return False, "DataFrame is empty"
        
        if len(df) < self.min_rows:
            return False, f"Dataset has fewer than {self.min_rows} rows"
        
        if len(df.columns) < self.min_cols:
            return False, f"Dataset has fewer than {self.min_cols} columns"
        
        null_cols = df.columns[df.isnull().all()].tolist()
        if len(null_cols) == len(df.columns):
            return False, "All columns are completely null"
        
        memory_mb = df.memory_usage(deep=True).sum() / (1024 * 1024)
        if memory_mb > self.max_size_mb:
            return False, f"Dataset size ({memory_mb:.1f} MB) exceeds limit ({self.max_size_mb} MB)"
        
        if len(df.columns) != len(set(df.columns)):
            return False, "Dataset contains duplicate column names"
        
        return True, "Dataset validation passed"
    
    def get_dataset_health_score(self, df):
        """Calculate overall dataset health score (0-100)"""
        health_metrics = {}
        
        # Completeness (40 points)
        missing_pct = (df.isnull().sum().sum() / (len(df) * len(df.columns))) * 100
        completeness_score = max(0, 40 - (missing_pct * 2))
        health_metrics['completeness'] = completeness_score
        
        # Consistency (30 points)
        consistency_issues = 0
        for col in df.columns:
            if df[col].dtype == 'object':
                unique_types = df[col].dropna().apply(type).nunique()
                if unique_types > 1:
                    consistency_issues += 1
        
        consistency_score = max(0, 30 - (consistency_issues * 5))
        health_metrics['consistency'] = consistency_score
        
        # Uniqueness (15 points)
        duplicate_pct = (df.duplicated().sum() / len(df)) * 100
        uniqueness_score = max(0, 15 - (duplicate_pct * 1.5))
        health_metrics['uniqueness'] = uniqueness_score
        
        # Validity (15 points)
        validity_issues = 0
        numeric_cols = df.select_dtypes(include=['number']).columns
        
        for col in numeric_cols:
            if np.isinf(df[col]).any():
                validity_issues += 1
        
        validity_score = max(0, 15 - (validity_issues * 3))
        health_metrics['validity'] = validity_score
        
        health_metrics['health_score'] = sum(health_metrics.values())
        
        return health_metrics

class ModelPersistence:
    """Handles model saving and loading"""
    
    def __init__(self, models_dir='models'):
        self.models_dir = models_dir
    
    def serialize_model(self, model):
        """Serialize model to bytes for download"""
        model_stream = io.BytesIO()
        pickle.dump(model, model_stream)
        model_stream.seek(0)
        return model_stream.getvalue()
    
    def deserialize_model(self, model_bytes):
        """Deserialize model from bytes"""
        model_stream = io.BytesIO(model_bytes)
        return pickle.load(model_stream)

class DataProfiler:
    """Provides detailed data profiling"""
    
    @staticmethod
    def profile_dataframe(df):
        """Generate comprehensive data profile"""
        profile = {
            'overview': {},
            'numeric_features': {},
            'categorical_features': {},
            'missing_values': {}
        }
        
        profile['overview'] = {
            'rows': len(df),
            'columns': len(df.columns),
            'memory_mb': df.memory_usage(deep=True).sum() / (1024 * 1024),
            'duplicates': df.duplicated().sum()
        }
        
        numeric_cols = df.select_dtypes(include=['number']).columns
        for col in numeric_cols:
            profile['numeric_features'][col] = {
                'mean': df[col].mean(),
                'median': df[col].median(),
                'std': df[col].std(),
                'min': df[col].min(),
                'max': df[col].max(),
                'missing': df[col].isnull().sum()
            }
        
        cat_cols = df.select_dtypes(include=['object', 'category']).columns
        for col in cat_cols:
            profile['categorical_features'][col] = {
                'unique_values': df[col].nunique(),
                'most_common': df[col].mode()[0] if len(df[col].mode()) > 0 else None,
                'missing': df[col].isnull().sum()
            }
        
        missing = df.isnull().sum()
        profile['missing_values'] = {
            'total': missing.sum(),
            'percentage': (missing.sum() / (len(df) * len(df.columns))) * 100,
            'by_column': missing[missing > 0].to_dict()
        }
        
        return profile

class ConfigManager:
    """Manages application configuration"""
    
    def __init__(self):
        self.config = {
            'random_state': 42,
            'max_features_plot': 15,
            'plot_dpi': 100,
            'test_size': 0.2,
            'cv_folds': 5,
            'use_smote': True
        }
    
    def get(self, key, default=None):
        return self.config.get(key, default)
    
    def set(self, key, value):
        self.config[key] = value
    
    def update(self, config_dict):
        self.config.update(config_dict)
'''
}

def create_all_modules():
    """Create all module files"""
    print("🚀 Creating AI Data Analysis Agent Module Files")
    print("=" * 60)
    
    created_files = []
    
    for filename, content in MODULES.items():
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(content)
            created_files.append(filename)
            print(f"✅ Created: {filename}")
        except Exception as e:
            print(f"❌ Failed to create {filename}: {str(e)}")
    
    print("\n" + "=" * 60)
    print(f"✨ Successfully created {len(created_files)} module files!")
    print("=" * 60)
    
    print("\n📋 Next Steps:")
    print("1. Run: python setup.py")
    print("2. Run: streamlit run integrated_app.py")
    print("3. Upload your data and start analyzing!")
    
    return created_files

if __name__ == "__main__":
    create_all_modules()
